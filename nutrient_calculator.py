import tkinter as tk
from tkinter import ttk, messagebox
import json
import os

class NutrientCalculator:
    def __init__(self):
        # Molecular weights of elements
        self.element_weights = {
            'N': 14.007,
            'O': 15.999,
            'H': 1.008,
            'P': 30.974,
            'K': 39.098,
            'Ca': 40.078,
            'Mg': 24.305,
            'S': 32.065,
            'Fe': 55.845,
            'Zn': 65.38,
            'Mn': 54.938,
            'B': 10.811,
            'Cu': 63.546,
            'Mo': 95.95
        }

        # Formula storage file
        self.formula_file = 'nutrient_formulas.json'
        self.formulas = self.load_formulas()

        # Database storage file
        self.database_file = 'nutrient_database.json'
        self.load_database()

        # Chemical compounds database with their formulas and molecular weights
        self.chemicals = {
            # Nitrate sources
            'Calcium Nitrate': {'formula': 'Ca(NO3)2·4H2O', 'mw': 236.15, 'nutrients': {'NO3': 2, 'Ca': 1}},
            'Potassium Nitrate': {'formula': 'KNO3', 'mw': 101.10, 'nutrients': {'NO3': 1, 'K': 1}},
            'Magnesium Nitrate': {'formula': 'Mg(NO3)2·6H2O', 'mw': 256.41, 'nutrients': {'NO3': 2, 'Mg': 1}},

            # Ammonium sources
            'Ammonium Sulfate': {'formula': '(NH4)2SO4', 'mw': 132.14, 'nutrients': {'NH4': 2, 'S': 1}},
            'Ammonium Nitrate': {'formula': 'NH4NO3', 'mw': 80.04, 'nutrients': {'NH4': 1, 'NO3': 1}},

            # Phosphorus sources
            'Monopotassium Phosphate (MKP)': {'formula': 'KH2PO4', 'mw': 136.09, 'nutrients': {'P': 1, 'K': 1}},
            'Monoammonium Phosphate (MAP)': {'formula': 'NH4H2PO4', 'mw': 115.03, 'nutrients': {'P': 1, 'NH4': 1}},
            'Phosphoric Acid (85%)': {'formula': 'H3PO4', 'mw': 98.00, 'nutrients': {'P': 1}, 'purity': 0.85},

            # Potassium sources
            'Potassium Sulfate': {'formula': 'K2SO4', 'mw': 174.26, 'nutrients': {'K': 2, 'S': 1}},
            'Potassium Chloride': {'formula': 'KCl', 'mw': 74.55, 'nutrients': {'K': 1}},

            # Calcium sources
            'Calcium Chloride': {'formula': 'CaCl2·2H2O', 'mw': 147.01, 'nutrients': {'Ca': 1}},

            # Magnesium sources
            'Magnesium Sulfate (Epsom Salt)': {'formula': 'MgSO4·7H2O', 'mw': 246.47, 'nutrients': {'Mg': 1, 'S': 1}},

            # Micronutrients
            'Iron EDTA (13%)': {'formula': 'FeEDTA', 'mw': 367.05, 'nutrients': {'Fe': 1}, 'purity': 0.13},
            'Iron DTPA (11%)': {'formula': 'FeDTPA', 'mw': 468.17, 'nutrients': {'Fe': 1}, 'purity': 0.11},
            'Zinc Sulfate': {'formula': 'ZnSO4·7H2O', 'mw': 287.54, 'nutrients': {'Zn': 1, 'S': 1}},
            'Manganese Sulfate': {'formula': 'MnSO4·H2O', 'mw': 169.01, 'nutrients': {'Mn': 1, 'S': 1}},
            'Boric Acid': {'formula': 'H3BO3', 'mw': 61.83, 'nutrients': {'B': 1}},
            'Copper Sulfate': {'formula': 'CuSO4·5H2O', 'mw': 249.68, 'nutrients': {'Cu': 1, 'S': 1}},
            'Ammonium Molybdate': {'formula': '(NH4)6Mo7O24·4H2O', 'mw': 1235.86, 'nutrients': {'Mo': 7, 'NH4': 6}},
        }

        # Nutrient forms mapping (NO3, NH4 are just forms of N)
        # The actual calculations use element_weights directly
        self.nutrient_to_element = {
            'NO3': 'N',  # Nitrate form of Nitrogen
            'NH4': 'N',  # Ammonium form of Nitrogen
            'P': 'P',
            'K': 'K',
            'Ca': 'Ca',
            'Mg': 'Mg',
            'S': 'S',
            'Fe': 'Fe',
            'Zn': 'Zn',
            'Mn': 'Mn',
            'B': 'B',
            'Cu': 'Cu',
            'Mo': 'Mo'
        }

        # Supported units
        self.units = ['mg/L (ppm)', 'g/L', '%']

        # Database password
        self.db_password = "3edcQWERTY"

    def convert_to_mg_per_L(self, value, unit):
        """Convert any unit to mg/L"""
        if unit == 'mg/L (ppm)':
            return value
        elif unit == 'g/L':
            return value * 1000
        elif unit == '%':
            return value * 10000  # 1% = 10000 mg/L
        return value

    def convert_from_mg_per_L(self, value_mg_per_L, target_unit):
        """Convert mg/L to target unit"""
        if target_unit == 'mg/L (ppm)':
            return value_mg_per_L
        elif target_unit == 'g/L':
            return value_mg_per_L / 1000
        elif target_unit == '%':
            return value_mg_per_L / 10000
        return value_mg_per_L

    def calculate_nutrient_from_chemical(self, chemical_name, amount_mg_per_L):
        """Calculate nutrient concentrations from chemical amount (mg/L)"""
        if chemical_name not in self.chemicals:
            return None

        chemical = self.chemicals[chemical_name]
        mw = chemical['mw']
        purity = chemical.get('purity', 1.0)

        # Calculate moles of chemical per liter
        moles = (amount_mg_per_L * purity) / mw

        nutrients = {}
        for nutrient_form, count in chemical['nutrients'].items():
            # Get the base element (NO3 -> N, NH4 -> N, etc.)
            element = self.nutrient_to_element[nutrient_form]
            # mg/L of element = moles * count * molecular weight of element
            nutrients[nutrient_form] = moles * count * self.element_weights[element]

        return nutrients

    def calculate_chemical_from_nutrient(self, chemical_name, target_nutrient, target_concentration_mg_per_L):
        """Calculate chemical amount needed to achieve target nutrient concentration"""
        if chemical_name not in self.chemicals:
            return None

        chemical = self.chemicals[chemical_name]

        # Check if chemical contains the target nutrient
        if target_nutrient not in chemical['nutrients']:
            return None

        mw = chemical['mw']
        purity = chemical.get('purity', 1.0)
        count = chemical['nutrients'][target_nutrient]

        # Get the base element (NO3 -> N, NH4 -> N, etc.)
        element = self.nutrient_to_element[target_nutrient]

        # Calculate required moles of element
        required_moles = target_concentration_mg_per_L / self.element_weights[element]

        # Calculate moles of chemical needed
        chemical_moles = required_moles / count

        # Calculate mg/L of chemical needed
        chemical_amount = (chemical_moles * mw) / purity

        return chemical_amount

    def load_formulas(self):
        """Load saved formulas from JSON file"""
        if os.path.exists(self.formula_file):
            try:
                with open(self.formula_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def save_formulas(self):
        """Save formulas to JSON file"""
        with open(self.formula_file, 'w', encoding='utf-8') as f:
            json.dump(self.formulas, f, ensure_ascii=False, indent=2)

    def add_formula(self, name, description, target_nutrients, chemicals_list):
        """Add a new formula
        target_nutrients: dict of {nutrient: concentration_mg_per_L}
        chemicals_list: list of {chemical_name: amount_mg_per_L}
        """
        self.formulas[name] = {
            'description': description,
            'target_nutrients': target_nutrients,
            'chemicals': chemicals_list
        }
        self.save_formulas()

    def delete_formula(self, name):
        """Delete a formula"""
        if name in self.formulas:
            del self.formulas[name]
            self.save_formulas()

    def get_formula(self, name):
        """Get a formula by name"""
        return self.formulas.get(name)

    def load_database(self):
        """Load chemical database from JSON file"""
        if os.path.exists(self.database_file):
            try:
                with open(self.database_file, 'r', encoding='utf-8') as f:
                    saved_db = json.load(f)
                    # Update chemicals with saved data
                    self.chemicals.update(saved_db.get('chemicals', {}))
                    # Update element weights if saved
                    if 'element_weights' in saved_db:
                        self.element_weights.update(saved_db['element_weights'])
            except Exception as e:
                print(f"Error loading database: {e}")

    def save_database(self):
        """Save chemical database to JSON file"""
        try:
            database = {
                'element_weights': self.element_weights,
                'chemicals': self.chemicals
            }
            with open(self.database_file, 'w', encoding='utf-8') as f:
                json.dump(database, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving database: {e}")

    def calculate_concentrated_solution(self, chemicals_list, stock_concentration_multiplier, final_volume_liters):
        """Calculate amounts needed for concentrated stock solution
        chemicals_list: list of {chemical_name: amount_mg_per_L} at working concentration
        stock_concentration_multiplier: e.g., 100 for 100x stock
        final_volume_liters: volume of stock solution to prepare
        """
        stock_amounts = {}
        for chemical_name, working_amount in chemicals_list.items():
            # Amount in grams for the stock solution
            amount_in_stock = (working_amount * stock_concentration_multiplier * final_volume_liters) / 1000
            stock_amounts[chemical_name] = amount_in_stock

        return stock_amounts


class NutrientCalculatorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("โปรแกรมคำนวณธาตุอาหาร - Nutrient Calculator")
        self.root.geometry("950x750")

        self.calculator = NutrientCalculator()

        # Font size variable
        self.font_size = tk.IntVar(value=10)

        # Create menu bar
        self.create_menu_bar()

        # Create notebook for tabs
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Create tabs
        self.create_chemical_to_nutrient_tab()
        self.create_nutrient_to_chemical_tab()
        self.create_formula_manager_tab()
        self.create_concentrated_solution_tab()
        self.create_database_tab()

    def create_menu_bar(self):
        """Create menu bar with font size and export options"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # View menu
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="มุมมอง", menu=view_menu)

        # Font size submenu
        font_menu = tk.Menu(view_menu, tearoff=0)
        view_menu.add_cascade(label="ขนาดตัวอักษร", menu=font_menu)

        for size in [8, 9, 10, 11, 12, 14, 16, 18]:
            font_menu.add_radiobutton(label=f"{size} pt",
                                     variable=self.font_size,
                                     value=size,
                                     command=self.update_font_size)

        # Export menu
        export_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="ส่งออก", menu=export_menu)
        export_menu.add_command(label="ส่งออกเป็น PDF", command=self.export_to_pdf)
        export_menu.add_command(label="ส่งออกเป็น DOCX", command=self.export_to_docx)
        export_menu.add_command(label="ส่งออกเป็น CSV", command=self.export_to_csv)

    def update_font_size(self):
        """Update font size for all text widgets"""
        size = self.font_size.get()
        font = ('Courier', size)

        # Update all text widgets
        for widget in [self.chem_to_nut_result, self.nut_to_chem_result,
                      self.formula_details, self.stock_result, self.db_text]:
            try:
                widget.configure(font=font)
            except:
                pass

    def create_chemical_to_nutrient_tab(self):
        """Tab for converting chemical amounts to nutrient concentrations"""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="สารเคมี → ธาตุอาหาร")

        # Title
        title = ttk.Label(frame, text="คำนวณความเข้มข้นธาตุอาหารจากปริมาณสารเคมี",
                         font=('Arial', 12, 'bold'))
        title.pack(pady=10)

        # Input frame
        input_frame = ttk.LabelFrame(frame, text="ข้อมูลสารเคมี", padding=10)
        input_frame.pack(fill='x', padx=20, pady=10)

        # Chemical selection
        ttk.Label(input_frame, text="เลือกสารเคมี:").grid(row=0, column=0, sticky='w', pady=5)
        self.chem_to_nut_chemical = ttk.Combobox(input_frame,
                                                  values=list(self.calculator.chemicals.keys()),
                                                  width=40, state='readonly')
        self.chem_to_nut_chemical.grid(row=0, column=1, pady=5, padx=10)
        self.chem_to_nut_chemical.bind('<<ComboboxSelected>>', self.update_chemical_info)

        # Chemical info display
        self.chemical_info_label = ttk.Label(input_frame, text="", foreground='blue')
        self.chemical_info_label.grid(row=1, column=1, sticky='w', pady=5)

        # Amount input
        ttk.Label(input_frame, text="ปริมาณสารเคมี:").grid(row=2, column=0, sticky='w', pady=5)
        self.chem_amount = ttk.Entry(input_frame, width=15)
        self.chem_amount.grid(row=2, column=1, sticky='w', pady=5, padx=10)

        # Unit selection
        ttk.Label(input_frame, text="หน่วย:").grid(row=2, column=2, sticky='w', pady=5)
        self.chem_to_nut_unit = ttk.Combobox(input_frame,
                                             values=self.calculator.units,
                                             width=12, state='readonly')
        self.chem_to_nut_unit.grid(row=2, column=3, sticky='w', pady=5, padx=10)
        self.chem_to_nut_unit.current(0)  # Default to mg/L

        # Calculate button
        calc_btn = ttk.Button(input_frame, text="คำนวณ", command=self.calculate_chem_to_nut)
        calc_btn.grid(row=3, column=1, pady=10)

        # Results frame
        result_frame = ttk.LabelFrame(frame, text="ผลลัพธ์ - ความเข้มข้นธาตุอาหาร (mg/L)", padding=10)
        result_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Create text widget for results
        self.chem_to_nut_result = tk.Text(result_frame, height=15, width=70, font=('Courier', 10))
        self.chem_to_nut_result.pack(fill='both', expand=True)

        # Scrollbar
        scrollbar = ttk.Scrollbar(result_frame, command=self.chem_to_nut_result.yview)
        scrollbar.pack(side='right', fill='y')
        self.chem_to_nut_result.config(yscrollcommand=scrollbar.set)

    def create_nutrient_to_chemical_tab(self):
        """Tab for converting nutrient concentrations to chemical amounts"""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="ธาตุอาหาร → สารเคมี")

        # Title
        title = ttk.Label(frame, text="คำนวณปริมาณสารเคมีจากความเข้มข้นธาตุอาหารที่ต้องการ",
                         font=('Arial', 12, 'bold'))
        title.pack(pady=10)

        # Input frame
        input_frame = ttk.LabelFrame(frame, text="ข้อมูลที่ต้องการ", padding=10)
        input_frame.pack(fill='x', padx=20, pady=10)

        # Target nutrient
        ttk.Label(input_frame, text="ธาตุอาหารเป้าหมาย:").grid(row=0, column=0, sticky='w', pady=5)
        # Create display names for nutrients
        nutrient_display = []
        for nutrient_form in self.calculator.nutrient_to_element.keys():
            if nutrient_form in ['NO3', 'NH4']:
                element = self.calculator.nutrient_to_element[nutrient_form]
                nutrient_display.append(f"{nutrient_form}-{element}")
            else:
                nutrient_display.append(nutrient_form)
        self.target_nutrient = ttk.Combobox(input_frame,
                                           values=nutrient_display,
                                           width=15, state='readonly')
        self.target_nutrient.grid(row=0, column=1, sticky='w', pady=5, padx=10)
        self.target_nutrient.bind('<<ComboboxSelected>>', self.update_available_chemicals)

        # Target concentration
        ttk.Label(input_frame, text="ความเข้มข้นที่ต้องการ:").grid(row=1, column=0, sticky='w', pady=5)
        self.target_concentration = ttk.Entry(input_frame, width=15)
        self.target_concentration.grid(row=1, column=1, sticky='w', pady=5, padx=10)

        # Unit selection for target concentration
        ttk.Label(input_frame, text="หน่วย:").grid(row=1, column=2, sticky='w', pady=5)
        self.nut_to_chem_unit = ttk.Combobox(input_frame,
                                             values=self.calculator.units,
                                             width=12, state='readonly')
        self.nut_to_chem_unit.grid(row=1, column=3, sticky='w', pady=5, padx=10)
        self.nut_to_chem_unit.current(0)  # Default to mg/L

        # Chemical selection
        ttk.Label(input_frame, text="เลือกสารเคมีที่จะใช้:").grid(row=2, column=0, sticky='w', pady=5)
        self.nut_to_chem_chemical = ttk.Combobox(input_frame, width=40, state='readonly')
        self.nut_to_chem_chemical.grid(row=2, column=1, pady=5, padx=10)
        self.nut_to_chem_chemical.bind('<<ComboboxSelected>>', self.update_chemical_info_nut_to_chem)

        # Chemical info display
        self.chemical_info_label2 = ttk.Label(input_frame, text="", foreground='blue')
        self.chemical_info_label2.grid(row=3, column=1, sticky='w', pady=5)

        # Calculate button
        calc_btn = ttk.Button(input_frame, text="คำนวณ", command=self.calculate_nut_to_chem)
        calc_btn.grid(row=4, column=1, pady=10)

        # Results frame
        result_frame = ttk.LabelFrame(frame, text="ผลลัพธ์", padding=10)
        result_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Create text widget for results
        self.nut_to_chem_result = tk.Text(result_frame, height=15, width=70, font=('Courier', 10))
        self.nut_to_chem_result.pack(fill='both', expand=True)

        # Scrollbar
        scrollbar = ttk.Scrollbar(result_frame, command=self.nut_to_chem_result.yview)
        scrollbar.pack(side='right', fill='y')
        self.nut_to_chem_result.config(yscrollcommand=scrollbar.set)

    def create_formula_manager_tab(self):
        """Tab for managing saved formulas"""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="จัดการสูตร")

        # Title
        title = ttk.Label(frame, text="จัดการสูตรธาตุอาหาร",
                         font=('Arial', 12, 'bold'))
        title.pack(pady=10)

        # Main container with two columns
        main_container = ttk.Frame(frame)
        main_container.pack(fill='both', expand=True, padx=20, pady=10)

        # Left side - Formula list
        left_frame = ttk.LabelFrame(main_container, text="สูตรที่บันทึกไว้", padding=10)
        left_frame.pack(side='left', fill='both', expand=True, padx=(0, 10))

        # Formula listbox
        list_frame = ttk.Frame(left_frame)
        list_frame.pack(fill='both', expand=True)

        self.formula_listbox = tk.Listbox(list_frame, height=15, font=('Arial', 10))
        self.formula_listbox.pack(side='left', fill='both', expand=True)
        self.formula_listbox.bind('<<ListboxSelect>>', self.on_formula_select)

        scrollbar = ttk.Scrollbar(list_frame, command=self.formula_listbox.yview)
        scrollbar.pack(side='right', fill='y')
        self.formula_listbox.config(yscrollcommand=scrollbar.set)

        # Buttons for formula management
        btn_frame = ttk.Frame(left_frame)
        btn_frame.pack(fill='x', pady=(10, 0))

        ttk.Button(btn_frame, text="โหลดสูตร", command=self.load_selected_formula).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="แก้ไข", command=self.edit_selected_formula).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="ลบสูตร", command=self.delete_selected_formula).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="รีเฟรช", command=self.refresh_formula_list).pack(side='left', padx=5)

        # Right side - Formula details and save new
        right_frame = ttk.LabelFrame(main_container, text="รายละเอียดสูตร", padding=10)
        right_frame.pack(side='right', fill='both', expand=True)

        # Formula details display
        self.formula_details = tk.Text(right_frame, height=20, width=50, font=('Courier', 9))
        self.formula_details.pack(fill='both', expand=True)

        details_scrollbar = ttk.Scrollbar(right_frame, command=self.formula_details.yview)
        details_scrollbar.pack(side='right', fill='y')
        self.formula_details.config(yscrollcommand=details_scrollbar.set)

        # Save current formula section
        save_frame = ttk.LabelFrame(frame, text="บันทึกสูตรใหม่", padding=10)
        save_frame.pack(fill='x', padx=20, pady=10)

        ttk.Label(save_frame, text="ชื่อสูตร:").grid(row=0, column=0, sticky='w', pady=5)
        self.new_formula_name = ttk.Entry(save_frame, width=30)
        self.new_formula_name.grid(row=0, column=1, sticky='w', pady=5, padx=10, columnspan=2)

        ttk.Label(save_frame, text="คำอธิบาย:").grid(row=1, column=0, sticky='w', pady=5)
        self.new_formula_desc = ttk.Entry(save_frame, width=50)
        self.new_formula_desc.grid(row=1, column=1, sticky='w', pady=5, padx=10, columnspan=2)

        # Chemical selector with dropdown
        ttk.Label(save_frame, text="เพิ่มสารเคมี:", font=('Arial', 9, 'bold')).grid(row=2, column=0, sticky='w', pady=10, columnspan=3)

        chem_input_frame = ttk.Frame(save_frame)
        chem_input_frame.grid(row=3, column=0, columnspan=3, sticky='ew', pady=5)

        ttk.Label(chem_input_frame, text="สารเคมี:").grid(row=0, column=0, sticky='w', padx=5)
        self.formula_chem_select = ttk.Combobox(chem_input_frame,
                                                values=sorted(self.calculator.chemicals.keys()),
                                                width=35, state='readonly')
        self.formula_chem_select.grid(row=0, column=1, padx=5)

        ttk.Label(chem_input_frame, text="ปริมาณ:").grid(row=0, column=2, sticky='w', padx=5)
        self.formula_chem_amount = ttk.Entry(chem_input_frame, width=15)
        self.formula_chem_amount.grid(row=0, column=3, padx=5)

        ttk.Label(chem_input_frame, text="หน่วย:").grid(row=0, column=4, sticky='w', padx=5)
        self.formula_chem_unit = ttk.Combobox(chem_input_frame,
                                              values=self.calculator.units,
                                              width=12, state='readonly')
        self.formula_chem_unit.grid(row=0, column=5, padx=5)
        self.formula_chem_unit.current(0)  # Default to mg/L

        ttk.Button(chem_input_frame, text="เพิ่ม",
                  command=self.add_chemical_to_formula).grid(row=0, column=6, padx=5)

        # Chemical list display
        ttk.Label(save_frame, text="สารเคมีในสูตร:").grid(row=4, column=0, sticky='nw', pady=5)

        chem_list_frame = ttk.Frame(save_frame)
        chem_list_frame.grid(row=4, column=1, columnspan=2, pady=5, sticky='ew')

        self.formula_chem_listbox = tk.Listbox(chem_list_frame, height=6, width=60, font=('Courier', 9))
        self.formula_chem_listbox.pack(side='left', fill='both', expand=True)

        chem_scroll = ttk.Scrollbar(chem_list_frame, command=self.formula_chem_listbox.yview)
        chem_scroll.pack(side='right', fill='y')
        self.formula_chem_listbox.config(yscrollcommand=chem_scroll.set)

        btn_frame2 = ttk.Frame(save_frame)
        btn_frame2.grid(row=5, column=1, columnspan=2, pady=5)

        ttk.Button(btn_frame2, text="ลบรายการที่เลือก",
                  command=self.remove_chemical_from_formula).pack(side='left', padx=5)
        ttk.Button(btn_frame2, text="ล้างทั้งหมด",
                  command=self.clear_formula_chemicals).pack(side='left', padx=5)

        ttk.Button(save_frame, text="บันทึกสูตร",
                  command=self.save_new_formula).grid(row=6, column=1, pady=10)

        # Store chemicals temporarily
        self.temp_formula_chemicals = {}

        # Initial refresh
        self.refresh_formula_list()

    def create_concentrated_solution_tab(self):
        """Tab for calculating concentrated stock solutions"""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="สารละลายเข้มข้น")

        # Title
        title = ttk.Label(frame, text="คำนวณการเตรียมสารละลายเข้มข้น (Stock Solution)",
                         font=('Arial', 12, 'bold'))
        title.pack(pady=10)

        # Input frame
        input_frame = ttk.LabelFrame(frame, text="ข้อมูลการเตรียมสารละลาย", padding=10)
        input_frame.pack(fill='x', padx=20, pady=10)

        # Load formula button
        ttk.Label(input_frame, text="โหลดสูตรที่บันทึก:").grid(row=0, column=0, sticky='w', pady=5)
        self.stock_formula_combo = ttk.Combobox(input_frame, width=30, state='readonly')
        self.stock_formula_combo.grid(row=0, column=1, sticky='w', pady=5, padx=10)
        ttk.Button(input_frame, text="โหลด",
                  command=self.load_formula_to_stock).grid(row=0, column=2, pady=5, padx=5)

        # Concentration multiplier
        ttk.Label(input_frame, text="ความเข้มข้น (เท่า):").grid(row=1, column=0, sticky='w', pady=5)
        self.stock_multiplier = ttk.Entry(input_frame, width=15)
        self.stock_multiplier.grid(row=1, column=1, sticky='w', pady=5, padx=10)
        self.stock_multiplier.insert(0, "100")
        ttk.Label(input_frame, text="(เช่น 100 = สารละลายเข้มข้น 100 เท่า)",
                 foreground='gray').grid(row=1, column=2, sticky='w', pady=5)

        # Final volume
        ttk.Label(input_frame, text="ปริมาตรที่ต้องการ (ลิตร):").grid(row=2, column=0, sticky='w', pady=5)
        self.stock_volume = ttk.Entry(input_frame, width=15)
        self.stock_volume.grid(row=2, column=1, sticky='w', pady=5, padx=10)
        self.stock_volume.insert(0, "1")

        # Chemical list input
        ttk.Label(input_frame, text="สารเคมี (JSON format):").grid(row=3, column=0, sticky='nw', pady=5)
        self.stock_chemicals = tk.Text(input_frame, height=5, width=50)
        self.stock_chemicals.grid(row=3, column=1, pady=5, padx=10, columnspan=2)

        example_text = '{"Calcium Nitrate": 945, "Potassium Nitrate": 810}'
        ttk.Label(input_frame, text=f"ตัวอย่าง (ความเข้มข้นที่ใช้งาน mg/L): {example_text}",
                 foreground='gray').grid(row=4, column=1, sticky='w', pady=0, padx=10, columnspan=2)

        # Calculate button
        ttk.Button(input_frame, text="คำนวณ",
                  command=self.calculate_stock_solution).grid(row=5, column=1, pady=10)

        # Results frame
        result_frame = ttk.LabelFrame(frame, text="ผลการคำนวณ", padding=10)
        result_frame.pack(fill='both', expand=True, padx=20, pady=10)

        self.stock_result = tk.Text(result_frame, height=15, width=70, font=('Courier', 10))
        self.stock_result.pack(fill='both', expand=True)

        scrollbar = ttk.Scrollbar(result_frame, command=self.stock_result.yview)
        scrollbar.pack(side='right', fill='y')
        self.stock_result.config(yscrollcommand=scrollbar.set)

        # Update formula combo
        self.update_stock_formula_combo()

    def create_database_tab(self):
        """Tab for viewing chemical database"""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="ฐานข้อมูลสารเคมี")

        # Title
        title = ttk.Label(frame, text="ฐานข้อมูลสารเคมีและน้ำหนักโมเลกุล",
                         font=('Arial', 12, 'bold'))
        title.pack(pady=10)

        # Password protection frame
        self.db_locked = True
        self.db_lock_frame = ttk.Frame(frame)
        self.db_lock_frame.pack(fill='both', expand=True, padx=20, pady=50)

        lock_inner = ttk.Frame(self.db_lock_frame)
        lock_inner.pack(expand=True)

        ttk.Label(lock_inner, text="🔒 ฐานข้อมูลถูกล็อค",
                 font=('Arial', 14, 'bold')).pack(pady=10)
        ttk.Label(lock_inner, text="กรุณาใส่รหัสผ่านเพื่อดูและแก้ไขฐานข้อมูล",
                 font=('Arial', 10)).pack(pady=5)

        pw_frame = ttk.Frame(lock_inner)
        pw_frame.pack(pady=20)

        ttk.Label(pw_frame, text="รหัสผ่าน:").pack(side='left', padx=5)
        self.db_password_entry = ttk.Entry(pw_frame, width=20, show='*')
        self.db_password_entry.pack(side='left', padx=5)
        self.db_password_entry.bind('<Return>', lambda e: self.unlock_database())

        ttk.Button(pw_frame, text="ปลดล็อค", command=self.unlock_database).pack(side='left', padx=5)

        # Database content frame (hidden initially)
        self.db_content_frame = ttk.Frame(frame)

        # Create text widget for database
        text_frame = ttk.Frame(self.db_content_frame)
        text_frame.pack(fill='both', expand=True, padx=20, pady=10)

        self.db_text = tk.Text(text_frame, height=20, width=80, font=('Courier', 9))
        self.db_text.pack(side='left', fill='both', expand=True)

        # Scrollbar
        scrollbar = ttk.Scrollbar(text_frame, command=self.db_text.yview)
        scrollbar.pack(side='right', fill='y')
        self.db_text.config(yscrollcommand=scrollbar.set)

        # Edit controls
        edit_frame = ttk.LabelFrame(self.db_content_frame, text="แก้ไขฐานข้อมูล", padding=10)
        edit_frame.pack(fill='x', padx=20, pady=10)

        btn_frame = ttk.Frame(edit_frame)
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text="แก้ไขน้ำหนักโมเลกุลธาตุ",
                  command=self.edit_element_weights).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="แก้ไขฐานข้อมูลสารเคมี (JSON)",
                  command=self.edit_database).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="ล็อคฐานข้อมูล",
                  command=self.lock_database).pack(side='left', padx=5)

    def display_database(self):
        """Display chemical database"""
        self.db_text.delete(1.0, tk.END)

        # Display element weights
        self.db_text.insert(tk.END, "=" * 80 + "\n")
        self.db_text.insert(tk.END, "น้ำหนักโมเลกุลของธาตุพื้นฐาน (Element Molecular Weights)\n")
        self.db_text.insert(tk.END, "=" * 80 + "\n\n")
        self.db_text.insert(tk.END, "หมายเหตุ: โปรแกรมใช้น้ำหนักโมเลกุลของธาตุพื้นฐานในการคำนวณ\n")
        self.db_text.insert(tk.END, "          NO3 และ NH4 เป็นเพียงรูปแบบของ N (ใช้น้ำหนัก N = 14.007)\n\n")

        for element, weight in self.calculator.element_weights.items():
            self.db_text.insert(tk.END, f"{element:<10} : {weight:>10.3f} g/mol\n")

        # Display chemicals
        self.db_text.insert(tk.END, "\n" + "=" * 80 + "\n")
        self.db_text.insert(tk.END, "ฐานข้อมูลสารเคมี (Chemical Database)\n")
        self.db_text.insert(tk.END, "=" * 80 + "\n\n")

        for name, info in self.calculator.chemicals.items():
            self.db_text.insert(tk.END, f"ชื่อ: {name}\n")
            self.db_text.insert(tk.END, f"  สูตร: {info['formula']}\n")
            self.db_text.insert(tk.END, f"  MW: {info['mw']:.2f} g/mol\n")
            if 'purity' in info:
                self.db_text.insert(tk.END, f"  ความบริสุทธิ์: {info['purity']*100:.0f}%\n")
            self.db_text.insert(tk.END, f"  ธาตุที่ให้: {', '.join(info['nutrients'].keys())}\n")
            self.db_text.insert(tk.END, "\n")

    def update_chemical_info(self, event=None):
        """Update chemical information display"""
        chemical_name = self.chem_to_nut_chemical.get()
        if chemical_name:
            info = self.calculator.chemicals[chemical_name]
            formula = info['formula']
            mw = info['mw']
            purity = info.get('purity', 1.0)
            purity_text = f", {purity*100:.0f}%" if purity != 1.0 else ""
            self.chemical_info_label.config(text=f"สูตร: {formula}, MW: {mw:.2f} g/mol{purity_text}")

    def update_chemical_info_nut_to_chem(self, event=None):
        """Update chemical information display in nutrient to chemical tab"""
        chemical_name = self.nut_to_chem_chemical.get()
        if chemical_name:
            info = self.calculator.chemicals[chemical_name]
            formula = info['formula']
            mw = info['mw']
            purity = info.get('purity', 1.0)
            purity_text = f", {purity*100:.0f}%" if purity != 1.0 else ""
            self.chemical_info_label2.config(text=f"สูตร: {formula}, MW: {mw:.2f} g/mol{purity_text}")

    def update_available_chemicals(self, event=None):
        """Update available chemicals based on selected nutrient"""
        nutrient_display = self.target_nutrient.get()
        if nutrient_display:
            # Convert display name back to nutrient form (NO3-N -> NO3)
            if '-' in nutrient_display:
                nutrient = nutrient_display.split('-')[0]
            else:
                nutrient = nutrient_display
            available = [name for name, info in self.calculator.chemicals.items()
                        if nutrient in info['nutrients']]
            self.nut_to_chem_chemical['values'] = available
            if available:
                self.nut_to_chem_chemical.current(0)
                self.update_chemical_info_nut_to_chem()

    def calculate_chem_to_nut(self):
        """Calculate nutrient concentrations from chemical amount"""
        try:
            chemical_name = self.chem_to_nut_chemical.get()
            amount = float(self.chem_amount.get())
            unit = self.chem_to_nut_unit.get()

            if not chemical_name:
                messagebox.showwarning("คำเตือน", "กรุณาเลือกสารเคมี")
                return

            # Convert to mg/L
            amount_mg_per_L = self.calculator.convert_to_mg_per_L(amount, unit)

            nutrients = self.calculator.calculate_nutrient_from_chemical(chemical_name, amount_mg_per_L)

            if nutrients:
                self.chem_to_nut_result.delete(1.0, tk.END)
                self.chem_to_nut_result.insert(tk.END, "=" * 60 + "\n")
                self.chem_to_nut_result.insert(tk.END, f"สารเคมี: {chemical_name}\n")
                self.chem_to_nut_result.insert(tk.END, f"ปริมาณ: {amount:.3f} {unit}\n")
                self.chem_to_nut_result.insert(tk.END, f"       = {amount_mg_per_L:.3f} mg/L\n")
                self.chem_to_nut_result.insert(tk.END, "=" * 60 + "\n\n")
                self.chem_to_nut_result.insert(tk.END, "ความเข้มข้นธาตุอาหารที่ได้:\n")
                self.chem_to_nut_result.insert(tk.END, "(หมายเหตุ: NO3-N และ NH4-N ใช้น้ำหนักโมเลกุลของ N = 14.007)\n\n")

                for nutrient_form, concentration in nutrients.items():
                    # Show form name (NO3, NH4) with element info
                    element = self.calculator.nutrient_to_element[nutrient_form]
                    if nutrient_form in ['NO3', 'NH4']:
                        display_name = f"{nutrient_form}-{element}"
                    else:
                        display_name = nutrient_form
                    self.chem_to_nut_result.insert(tk.END,
                        f"  {display_name:<10} : {concentration:>10.3f} mg/L\n")

                self.chem_to_nut_result.insert(tk.END, "\n" + "=" * 60 + "\n")

        except ValueError:
            messagebox.showerror("ข้อผิดพลาด", "กรุณาใส่ตัวเลขที่ถูกต้อง")

    def calculate_nut_to_chem(self):
        """Calculate chemical amount from target nutrient concentration"""
        try:
            nutrient_display = self.target_nutrient.get()
            concentration = float(self.target_concentration.get())
            unit = self.nut_to_chem_unit.get()
            chemical_name = self.nut_to_chem_chemical.get()

            if not nutrient_display or not chemical_name:
                messagebox.showwarning("คำเตือน", "กรุณาเลือกธาตุอาหารและสารเคมี")
                return

            # Convert display name back to nutrient form (NO3-N -> NO3)
            if '-' in nutrient_display:
                nutrient = nutrient_display.split('-')[0]
            else:
                nutrient = nutrient_display

            # Convert to mg/L
            concentration_mg_per_L = self.calculator.convert_to_mg_per_L(concentration, unit)

            chemical_amount = self.calculator.calculate_chemical_from_nutrient(
                chemical_name, nutrient, concentration_mg_per_L)

            if chemical_amount:
                # Calculate all nutrients from this chemical amount
                all_nutrients = self.calculator.calculate_nutrient_from_chemical(
                    chemical_name, chemical_amount)

                self.nut_to_chem_result.delete(1.0, tk.END)
                self.nut_to_chem_result.insert(tk.END, "=" * 60 + "\n")
                self.nut_to_chem_result.insert(tk.END, f"ธาตุอาหารเป้าหมาย: {nutrient_display}\n")
                self.nut_to_chem_result.insert(tk.END, f"ความเข้มข้นที่ต้องการ: {concentration:.3f} {unit}\n")
                self.nut_to_chem_result.insert(tk.END, f"                    = {concentration_mg_per_L:.3f} mg/L\n")
                self.nut_to_chem_result.insert(tk.END, "=" * 60 + "\n\n")

                self.nut_to_chem_result.insert(tk.END, f"ปริมาณสารเคมีที่ต้องใช้:\n\n")
                self.nut_to_chem_result.insert(tk.END,
                    f"  {chemical_name}\n")
                self.nut_to_chem_result.insert(tk.END,
                    f"  จำนวน: {chemical_amount:.3f} mg/L\n")
                self.nut_to_chem_result.insert(tk.END,
                    f"         {chemical_amount/1000:.3f} g/L\n")
                self.nut_to_chem_result.insert(tk.END,
                    f"         {chemical_amount/10000:.4f} %\n\n")

                self.nut_to_chem_result.insert(tk.END, "=" * 60 + "\n")
                self.nut_to_chem_result.insert(tk.END, "ธาตุอาหารทั้งหมดที่ได้:\n")
                self.nut_to_chem_result.insert(tk.END, "(หมายเหตุ: NO3-N และ NH4-N ใช้น้ำหนักโมเลกุลของ N = 14.007)\n\n")

                for nutrient_form, conc in all_nutrients.items():
                    marker = " ← เป้าหมาย" if nutrient_form == nutrient else ""
                    # Show form name (NO3, NH4) with element info
                    element = self.calculator.nutrient_to_element[nutrient_form]
                    if nutrient_form in ['NO3', 'NH4']:
                        display_name = f"{nutrient_form}-{element}"
                    else:
                        display_name = nutrient_form
                    self.nut_to_chem_result.insert(tk.END,
                        f"  {display_name:<10} : {conc:>10.3f} mg/L{marker}\n")

                self.nut_to_chem_result.insert(tk.END, "\n" + "=" * 60 + "\n")
            else:
                messagebox.showerror("ข้อผิดพลาด",
                    f"สารเคมี {chemical_name} ไม่มีธาตุ {nutrient}")

        except ValueError:
            messagebox.showerror("ข้อผิดพลาด", "กรุณาใส่ตัวเลขที่ถูกต้อง")

    # Formula management methods
    def refresh_formula_list(self):
        """Refresh the formula listbox"""
        self.formula_listbox.delete(0, tk.END)
        for name in self.calculator.formulas.keys():
            self.formula_listbox.insert(tk.END, name)

    def on_formula_select(self, event):
        """Display formula details when selected"""
        selection = self.formula_listbox.curselection()
        if selection:
            formula_name = self.formula_listbox.get(selection[0])
            formula = self.calculator.get_formula(formula_name)

            self.formula_details.delete(1.0, tk.END)
            self.formula_details.insert(tk.END, "=" * 50 + "\n")
            self.formula_details.insert(tk.END, f"ชื่อสูตร: {formula_name}\n")
            self.formula_details.insert(tk.END, f"คำอธิบาย: {formula.get('description', '-')}\n")
            self.formula_details.insert(tk.END, "=" * 50 + "\n\n")

            # Display chemicals
            self.formula_details.insert(tk.END, "สารเคมีและปริมาณ (mg/L):\n\n")
            for chem, amount in formula.get('chemicals', {}).items():
                self.formula_details.insert(tk.END, f"  {chem:<35} : {amount:>10.2f}\n")

            # Calculate total nutrients
            self.formula_details.insert(tk.END, "\n" + "=" * 50 + "\n")
            self.formula_details.insert(tk.END, "ธาตุอาหารทั้งหมดที่ได้:\n\n")

            total_nutrients = {}
            for chem_name, amount in formula.get('chemicals', {}).items():
                nutrients = self.calculator.calculate_nutrient_from_chemical(chem_name, amount)
                if nutrients:
                    for nutrient, conc in nutrients.items():
                        total_nutrients[nutrient] = total_nutrients.get(nutrient, 0) + conc

            for nutrient, conc in total_nutrients.items():
                self.formula_details.insert(tk.END, f"  {nutrient:<10} : {conc:>10.3f} mg/L\n")

    def load_selected_formula(self):
        """Load selected formula to clipboard or show it"""
        selection = self.formula_listbox.curselection()
        if selection:
            formula_name = self.formula_listbox.get(selection[0])
            formula = self.calculator.get_formula(formula_name)
            messagebox.showinfo("สูตรที่เลือก",
                              f"สูตร: {formula_name}\n\nสามารถดูรายละเอียดทางด้านขวา")
        else:
            messagebox.showwarning("คำเตือน", "กรุณาเลือกสูตรก่อน")

    def delete_selected_formula(self):
        """Delete selected formula"""
        selection = self.formula_listbox.curselection()
        if selection:
            formula_name = self.formula_listbox.get(selection[0])
            result = messagebox.askyesno("ยืนยันการลบ",
                                        f"คุณต้องการลบสูตร '{formula_name}' ใช่หรือไม่?")
            if result:
                self.calculator.delete_formula(formula_name)
                self.refresh_formula_list()
                self.update_stock_formula_combo()
                self.formula_details.delete(1.0, tk.END)
                messagebox.showinfo("สำเร็จ", "ลบสูตรเรียบร้อยแล้ว")
        else:
            messagebox.showwarning("คำเตือน", "กรุณาเลือกสูตรที่จะลบ")

    def edit_selected_formula(self):
        """Edit selected formula"""
        selection = self.formula_listbox.curselection()
        if not selection:
            messagebox.showwarning("คำเตือน", "กรุณาเลือกสูตรที่จะแก้ไข")
            return

        formula_name = self.formula_listbox.get(selection[0])
        formula = self.calculator.get_formula(formula_name)

        if not formula:
            messagebox.showerror("ข้อผิดพลาด", "ไม่พบสูตร")
            return

        # Create edit window
        edit_window = tk.Toplevel(self.root)
        edit_window.title(f"แก้ไขสูตร: {formula_name}")
        edit_window.geometry("800x700")

        # Title
        ttk.Label(edit_window, text=f"แก้ไขสูตร: {formula_name}",
                 font=('Arial', 14, 'bold')).pack(pady=10)

        # Form frame
        form_frame = ttk.LabelFrame(edit_window, text="ข้อมูลสูตร", padding=10)
        form_frame.pack(fill='x', padx=20, pady=10)

        # Name
        ttk.Label(form_frame, text="ชื่อสูตร:").grid(row=0, column=0, sticky='w', pady=5)
        edit_name = ttk.Entry(form_frame, width=40)
        edit_name.grid(row=0, column=1, sticky='w', pady=5, padx=10, columnspan=2)
        edit_name.insert(0, formula_name)

        # Description
        ttk.Label(form_frame, text="คำอธิบาย:").grid(row=1, column=0, sticky='w', pady=5)
        edit_desc = ttk.Entry(form_frame, width=60)
        edit_desc.grid(row=1, column=1, sticky='w', pady=5, padx=10, columnspan=3)
        edit_desc.insert(0, formula.get('description', ''))

        # Chemical editor frame
        chem_frame = ttk.LabelFrame(edit_window, text="สารเคมีในสูตร", padding=10)
        chem_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Add chemical controls
        add_frame = ttk.Frame(chem_frame)
        add_frame.pack(fill='x', pady=5)

        ttk.Label(add_frame, text="สารเคมี:").grid(row=0, column=0, sticky='w', padx=5)
        edit_chem_select = ttk.Combobox(add_frame,
                                        values=sorted(self.calculator.chemicals.keys()),
                                        width=30, state='readonly')
        edit_chem_select.grid(row=0, column=1, padx=5)

        ttk.Label(add_frame, text="ปริมาณ:").grid(row=0, column=2, sticky='w', padx=5)
        edit_chem_amount = ttk.Entry(add_frame, width=12)
        edit_chem_amount.grid(row=0, column=3, padx=5)

        ttk.Label(add_frame, text="หน่วย:").grid(row=0, column=4, sticky='w', padx=5)
        edit_chem_unit = ttk.Combobox(add_frame,
                                      values=self.calculator.units,
                                      width=12, state='readonly')
        edit_chem_unit.grid(row=0, column=5, padx=5)
        edit_chem_unit.current(0)

        # Temporary storage for edited chemicals
        edit_temp_chemicals = formula.get('chemicals', {}).copy()

        # Chemical listbox
        list_frame = ttk.Frame(chem_frame)
        list_frame.pack(fill='both', expand=True, pady=10)

        edit_chem_listbox = tk.Listbox(list_frame, height=10, width=70, font=('Courier', 9))
        edit_chem_listbox.pack(side='left', fill='both', expand=True)

        chem_scroll = ttk.Scrollbar(list_frame, command=edit_chem_listbox.yview)
        chem_scroll.pack(side='right', fill='y')
        edit_chem_listbox.config(yscrollcommand=chem_scroll.set)

        def update_chem_listbox():
            """Update chemical listbox"""
            edit_chem_listbox.delete(0, tk.END)
            for chem, amount in edit_temp_chemicals.items():
                display_text = f"{chem:<35} : {amount:>10.2f} mg/L"
                edit_chem_listbox.insert(tk.END, display_text)

        def add_chemical():
            """Add chemical to edited formula"""
            try:
                chemical = edit_chem_select.get()
                amount_str = edit_chem_amount.get().strip()
                unit = edit_chem_unit.get()

                if not chemical or not amount_str:
                    messagebox.showwarning("คำเตือน", "กรุณาเลือกสารเคมีและใส่ปริมาณ")
                    return

                amount = float(amount_str)
                amount_mg_per_L = self.calculator.convert_to_mg_per_L(amount, unit)

                edit_temp_chemicals[chemical] = amount_mg_per_L
                update_chem_listbox()
                edit_chem_amount.delete(0, tk.END)

            except ValueError:
                messagebox.showerror("ข้อผิดพลาด", "กรุณาใส่ตัวเลขที่ถูกต้อง")

        def remove_chemical():
            """Remove selected chemical"""
            selection = edit_chem_listbox.curselection()
            if selection:
                idx = selection[0]
                chemical_name = list(edit_temp_chemicals.keys())[idx]
                del edit_temp_chemicals[chemical_name]
                update_chem_listbox()
            else:
                messagebox.showwarning("คำเตือน", "กรุณาเลือกรายการที่จะลบ")

        def save_changes():
            """Save edited formula"""
            try:
                new_name = edit_name.get().strip()
                new_desc = edit_desc.get().strip()

                if not new_name:
                    messagebox.showwarning("คำเตือน", "กรุณาใส่ชื่อสูตร")
                    return

                if not edit_temp_chemicals:
                    messagebox.showwarning("คำเตือน", "กรุณาเพิ่มสารเคมีอย่างน้อย 1 รายการ")
                    return

                # Calculate nutrients
                target_nutrients = {}
                for chem_name, amount in edit_temp_chemicals.items():
                    nutrients = self.calculator.calculate_nutrient_from_chemical(chem_name, amount)
                    if nutrients:
                        for nutrient, conc in nutrients.items():
                            target_nutrients[nutrient] = target_nutrients.get(nutrient, 0) + conc

                # Delete old formula if name changed
                if new_name != formula_name:
                    self.calculator.delete_formula(formula_name)

                # Save formula
                self.calculator.add_formula(new_name, new_desc, target_nutrients, edit_temp_chemicals)

                # Refresh
                self.refresh_formula_list()
                self.update_stock_formula_combo()
                self.formula_details.delete(1.0, tk.END)

                messagebox.showinfo("สำเร็จ", f"บันทึกการแก้ไขสูตร '{new_name}' เรียบร้อยแล้ว")
                edit_window.destroy()

            except Exception as e:
                messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

        # Button frame for chemicals
        chem_btn_frame = ttk.Frame(chem_frame)
        chem_btn_frame.pack(fill='x', pady=5)

        ttk.Button(chem_btn_frame, text="เพิ่ม", command=add_chemical).pack(side='left', padx=5)
        ttk.Button(chem_btn_frame, text="ลบรายการที่เลือก", command=remove_chemical).pack(side='left', padx=5)

        # Initial update
        update_chem_listbox()

        # Save/Cancel buttons
        btn_frame = ttk.Frame(edit_window)
        btn_frame.pack(fill='x', padx=20, pady=10)

        ttk.Button(btn_frame, text="บันทึกการแก้ไข", command=save_changes).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="ยกเลิก", command=edit_window.destroy).pack(side='left', padx=5)

    def add_chemical_to_formula(self):
        """Add chemical to formula list"""
        try:
            chemical = self.formula_chem_select.get()
            amount_str = self.formula_chem_amount.get().strip()
            unit = self.formula_chem_unit.get()

            if not chemical:
                messagebox.showwarning("คำเตือน", "กรุณาเลือกสารเคมี")
                return

            if not amount_str:
                messagebox.showwarning("คำเตือน", "กรุณาใส่ปริมาณ")
                return

            amount = float(amount_str)

            # Convert to mg/L
            amount_mg_per_L = self.calculator.convert_to_mg_per_L(amount, unit)

            # Add to temp storage
            self.temp_formula_chemicals[chemical] = amount_mg_per_L

            # Update listbox
            self.update_formula_chemical_listbox()

            # Clear inputs
            self.formula_chem_amount.delete(0, tk.END)

        except ValueError:
            messagebox.showerror("ข้อผิดพลาด", "กรุณาใส่ตัวเลขที่ถูกต้อง")

    def remove_chemical_from_formula(self):
        """Remove selected chemical from formula"""
        selection = self.formula_chem_listbox.curselection()
        if selection:
            idx = selection[0]
            chemical_name = list(self.temp_formula_chemicals.keys())[idx]
            del self.temp_formula_chemicals[chemical_name]
            self.update_formula_chemical_listbox()
        else:
            messagebox.showwarning("คำเตือน", "กรุณาเลือกรายการที่จะลบ")

    def clear_formula_chemicals(self):
        """Clear all chemicals from formula"""
        self.temp_formula_chemicals = {}
        self.update_formula_chemical_listbox()

    def update_formula_chemical_listbox(self):
        """Update the chemical listbox display"""
        self.formula_chem_listbox.delete(0, tk.END)
        for chemical, amount in self.temp_formula_chemicals.items():
            display_text = f"{chemical:<35} : {amount:>10.2f} mg/L"
            self.formula_chem_listbox.insert(tk.END, display_text)

    def save_new_formula(self):
        """Save new formula"""
        try:
            name = self.new_formula_name.get().strip()
            description = self.new_formula_desc.get().strip()

            if not name:
                messagebox.showwarning("คำเตือน", "กรุณาใส่ชื่อสูตร")
                return

            if not self.temp_formula_chemicals:
                messagebox.showwarning("คำเตือน", "กรุณาเพิ่มสารเคมีอย่างน้อย 1 รายการ")
                return

            # Calculate target nutrients
            target_nutrients = {}
            for chem_name, amount in self.temp_formula_chemicals.items():
                nutrients = self.calculator.calculate_nutrient_from_chemical(chem_name, amount)
                if nutrients:
                    for nutrient, conc in nutrients.items():
                        target_nutrients[nutrient] = target_nutrients.get(nutrient, 0) + conc

            # Save formula
            self.calculator.add_formula(name, description, target_nutrients, self.temp_formula_chemicals)

            # Clear input fields
            self.new_formula_name.delete(0, tk.END)
            self.new_formula_desc.delete(0, tk.END)
            self.temp_formula_chemicals = {}
            self.update_formula_chemical_listbox()

            # Refresh list
            self.refresh_formula_list()
            self.update_stock_formula_combo()

            messagebox.showinfo("สำเร็จ", f"บันทึกสูตร '{name}' เรียบร้อยแล้ว")

        except Exception as e:
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    # Concentrated solution methods
    def update_stock_formula_combo(self):
        """Update the formula combo in stock solution tab"""
        formulas = list(self.calculator.formulas.keys())
        self.stock_formula_combo['values'] = formulas

    def load_formula_to_stock(self):
        """Load selected formula to stock solution calculator"""
        formula_name = self.stock_formula_combo.get()
        if formula_name:
            formula = self.calculator.get_formula(formula_name)
            if formula:
                chemicals_json = json.dumps(formula['chemicals'], ensure_ascii=False, indent=2)
                self.stock_chemicals.delete(1.0, tk.END)
                self.stock_chemicals.insert(1.0, chemicals_json)
                messagebox.showinfo("สำเร็จ", f"โหลดสูตร '{formula_name}' เรียบร้อยแล้ว")
        else:
            messagebox.showwarning("คำเตือน", "กรุณาเลือกสูตรก่อน")

    def calculate_stock_solution(self):
        """Calculate concentrated stock solution"""
        try:
            multiplier = float(self.stock_multiplier.get())
            volume = float(self.stock_volume.get())
            chemicals_json = self.stock_chemicals.get(1.0, tk.END).strip()

            if not chemicals_json:
                messagebox.showwarning("คำเตือน", "กรุณาใส่ข้อมูลสารเคมี")
                return

            # Parse JSON
            chemicals = json.loads(chemicals_json)

            # Calculate stock amounts
            stock_amounts = self.calculator.calculate_concentrated_solution(
                chemicals, multiplier, volume)

            # Display results
            self.stock_result.delete(1.0, tk.END)
            self.stock_result.insert(tk.END, "=" * 60 + "\n")
            self.stock_result.insert(tk.END, f"ความเข้มข้นของสารละลาย: {multiplier:.0f}x\n")
            self.stock_result.insert(tk.END, f"ปริมาตรที่เตรียม: {volume:.2f} ลิตร\n")
            self.stock_result.insert(tk.END, "=" * 60 + "\n\n")

            self.stock_result.insert(tk.END, "ปริมาณสารเคมีที่ต้องชั่ง:\n\n")

            total_weight = 0
            for chem_name, amount_g in stock_amounts.items():
                self.stock_result.insert(tk.END, f"  {chem_name:<35} : {amount_g:>10.3f} g\n")
                total_weight += amount_g

            self.stock_result.insert(tk.END, f"\n{'รวมน้ำหนักทั้งหมด':<35} : {total_weight:>10.3f} g\n")

            self.stock_result.insert(tk.END, "\n" + "=" * 60 + "\n")
            self.stock_result.insert(tk.END, "วิธีเตรียม:\n")
            self.stock_result.insert(tk.END, f"1. ชั่งสารเคมีตามปริมาณที่ระบุข้างต้น\n")
            self.stock_result.insert(tk.END, f"2. ละลายในน้ำให้ได้ปริมาตร {volume:.2f} ลิตร\n")
            self.stock_result.insert(tk.END, f"3. เมื่อใช้งาน เจือจาง {multiplier:.0f} เท่า ")
            self.stock_result.insert(tk.END, f"(เช่น {1000/multiplier:.1f} ml ต่อน้ำ 1 ลิตร)\n")
            self.stock_result.insert(tk.END, "=" * 60 + "\n")

            # Show nutrient concentration at working dilution
            self.stock_result.insert(tk.END, "\nความเข้มข้นธาตุอาหารเมื่อเจือจาง:\n\n")
            total_nutrients = {}
            for chem_name, amount in chemicals.items():
                nutrients = self.calculator.calculate_nutrient_from_chemical(chem_name, amount)
                if nutrients:
                    for nutrient, conc in nutrients.items():
                        total_nutrients[nutrient] = total_nutrients.get(nutrient, 0) + conc

            for nutrient, conc in total_nutrients.items():
                self.stock_result.insert(tk.END, f"  {nutrient:<10} : {conc:>10.3f} mg/L\n")

        except json.JSONDecodeError:
            messagebox.showerror("ข้อผิดพลาด", "รูปแบบ JSON ไม่ถูกต้อง")
        except ValueError:
            messagebox.showerror("ข้อผิดพลาด", "กรุณาใส่ตัวเลขที่ถูกต้อง")
        except Exception as e:
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    # Database management methods
    def unlock_database(self):
        """Unlock database with password"""
        password = self.db_password_entry.get()
        if password == self.calculator.db_password:
            self.db_locked = False
            self.db_lock_frame.pack_forget()
            self.db_content_frame.pack(fill='both', expand=True, padx=20, pady=10)
            self.display_database()
            messagebox.showinfo("สำเร็จ", "ปลดล็อคฐานข้อมูลเรียบร้อยแล้ว")
        else:
            messagebox.showerror("ข้อผิดพลาด", "รหัสผ่านไม่ถูกต้อง")
            self.db_password_entry.delete(0, tk.END)

    def lock_database(self):
        """Lock database"""
        self.db_locked = True
        self.db_content_frame.pack_forget()
        self.db_lock_frame.pack(fill='both', expand=True, padx=20, pady=50)
        self.db_password_entry.delete(0, tk.END)
        messagebox.showinfo("สำเร็จ", "ล็อคฐานข้อมูลเรียบร้อยแล้ว")

    def edit_database(self):
        """Open database editor window"""
        if self.db_locked:
            messagebox.showwarning("คำเตือน", "กรุณาปลดล็อคฐานข้อมูลก่อน")
            return

        # Create edit window
        edit_window = tk.Toplevel(self.root)
        edit_window.title("แก้ไขฐานข้อมูลสารเคมี")
        edit_window.geometry("700x600")

        ttk.Label(edit_window, text="แก้ไขฐานข้อมูลสารเคมี (JSON Format)",
                 font=('Arial', 12, 'bold')).pack(pady=10)

        ttk.Label(edit_window, text="คำเตือน: กรุณาตรวจสอบรูปแบบ JSON ให้ถูกต้องก่อนบันทึก",
                 foreground='red').pack(pady=5)

        # Text editor
        text_frame = ttk.Frame(edit_window)
        text_frame.pack(fill='both', expand=True, padx=20, pady=10)

        edit_text = tk.Text(text_frame, height=25, width=80, font=('Courier', 9))
        edit_text.pack(side='left', fill='both', expand=True)

        scrollbar = ttk.Scrollbar(text_frame, command=edit_text.yview)
        scrollbar.pack(side='right', fill='y')
        edit_text.config(yscrollcommand=scrollbar.set)

        # Load current database
        db_json = json.dumps(self.calculator.chemicals, ensure_ascii=False, indent=2)
        edit_text.insert(1.0, db_json)

        # Buttons
        btn_frame = ttk.Frame(edit_window)
        btn_frame.pack(fill='x', padx=20, pady=10)

        def save_changes():
            try:
                new_db = json.loads(edit_text.get(1.0, tk.END))
                self.calculator.chemicals = new_db
                self.calculator.save_database()  # Save to file
                self.display_database()
                messagebox.showinfo("สำเร็จ", "บันทึกการเปลี่ยนแปลงเรียบร้อยแล้ว")
                edit_window.destroy()
            except json.JSONDecodeError as e:
                messagebox.showerror("ข้อผิดพลาด", f"รูปแบบ JSON ไม่ถูกต้อง:\n{str(e)}")

        ttk.Button(btn_frame, text="บันทึก", command=save_changes).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="ยกเลิก", command=edit_window.destroy).pack(side='left', padx=5)

    def edit_element_weights(self):
        """Open element weights editor window"""
        if self.db_locked:
            messagebox.showwarning("คำเตือน", "กรุณาปลดล็อคฐานข้อมูลก่อน")
            return

        # Create edit window
        edit_window = tk.Toplevel(self.root)
        edit_window.title("แก้ไขน้ำหนักโมเลกุลธาตุ")
        edit_window.geometry("600x650")

        ttk.Label(edit_window, text="แก้ไขน้ำหนักโมเลกุลของธาตุ (Element Molecular Weights)",
                 font=('Arial', 12, 'bold')).pack(pady=10)

        ttk.Label(edit_window, text="หน่วย: g/mol",
                 foreground='blue').pack(pady=5)

        # Scrollable frame for elements
        canvas = tk.Canvas(edit_window)
        scrollbar = ttk.Scrollbar(edit_window, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Create input fields for each element
        element_entries = {}

        input_frame = ttk.LabelFrame(scrollable_frame, text="น้ำหนักโมเลกุลของธาตุพื้นฐาน", padding=15)
        input_frame.pack(fill='both', expand=True, padx=20, pady=10)

        row = 0
        for element, weight in self.calculator.element_weights.items():
            # Element name
            ttk.Label(input_frame, text=f"{element}:",
                     font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky='w', pady=5, padx=10)

            # Current value
            entry = ttk.Entry(input_frame, width=15, font=('Arial', 10))
            entry.insert(0, str(weight))
            entry.grid(row=row, column=1, pady=5, padx=10)

            # Unit label
            ttk.Label(input_frame, text="g/mol",
                     foreground='gray').grid(row=row, column=2, sticky='w', pady=5)

            element_entries[element] = entry
            row += 1

        canvas.pack(side="left", fill="both", expand=True, padx=20, pady=10)
        scrollbar.pack(side="right", fill="y")

        # Show affected ions
        info_frame = ttk.LabelFrame(edit_window, text="ธาตุอาหารที่จะได้รับผลกระทบ", padding=10)
        info_frame.pack(fill='x', padx=20, pady=10)

        info_text = """
        การแก้ไขน้ำหนักโมเลกุลของธาตุจะส่งผลกระทบต่อ:

        • N, O, H → NO₃ (Nitrate), NH₄ (Ammonium)
        • P → Phosphorus
        • K → Potassium
        • Ca → Calcium
        • Mg → Magnesium
        • S → Sulfur
        • Fe, Zn, Mn, B, Cu → Micronutrients

        ระบบจะคำนวณน้ำหนักโมเลกุลของธาตุอาหารอัตโนมัติ
        """

        ttk.Label(info_frame, text=info_text,
                 font=('Arial', 9), justify='left').pack()

        # Buttons
        btn_frame = ttk.Frame(edit_window)
        btn_frame.pack(fill='x', padx=20, pady=10)

        def save_element_weights():
            try:
                # Validate and save
                new_weights = {}
                for element, entry in element_entries.items():
                    try:
                        weight = float(entry.get())
                        if weight <= 0:
                            raise ValueError(f"น้ำหนักโมเลกุลต้องมากกว่า 0")
                        new_weights[element] = weight
                    except ValueError as e:
                        messagebox.showerror("ข้อผิดพลาด",
                                           f"ข้อมูลไม่ถูกต้องสำหรับธาตุ {element}: {str(e)}")
                        return

                # Update element weights
                self.calculator.element_weights = new_weights

                # Save to file
                self.calculator.save_database()

                # Update display
                self.display_database()

                messagebox.showinfo("สำเร็จ",
                                  "บันทึกน้ำหนักโมเลกุลธาตุเรียบร้อยแล้ว\n" +
                                  "โปรแกรมใช้น้ำหนักโมเลกุลของธาตุในการคำนวณโดยตรง")
                edit_window.destroy()

            except Exception as e:
                messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

        def reset_to_default():
            """Reset to default IUPAC values"""
            result = messagebox.askyesno("ยืนยัน",
                                        "คุณต้องการรีเซ็ตเป็นค่ามาตรฐาน IUPAC ใช่หรือไม่?")
            if result:
                default_weights = {
                    'N': 14.007, 'O': 15.999, 'H': 1.008,
                    'P': 30.974, 'K': 39.098, 'Ca': 40.078,
                    'Mg': 24.305, 'S': 32.065, 'Fe': 55.845,
                    'Zn': 65.38, 'Mn': 54.938, 'B': 10.811, 'Cu': 63.546,
                    'Mo': 95.95
                }
                for element, weight in default_weights.items():
                    element_entries[element].delete(0, tk.END)
                    element_entries[element].insert(0, str(weight))

        ttk.Button(btn_frame, text="บันทึก",
                  command=save_element_weights).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="รีเซ็ตค่ามาตรฐาน",
                  command=reset_to_default).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="ยกเลิก",
                  command=edit_window.destroy).pack(side='left', padx=5)

    # Export functions
    def get_current_tab_content(self):
        """Get content from current active tab"""
        current_tab = self.notebook.select()
        tab_index = self.notebook.index(current_tab)

        content = {
            'title': '',
            'data': []
        }

        if tab_index == 0:  # Chemical to Nutrient
            content['title'] = 'คำนวณธาตุอาหารจากสารเคมี'
            content['data'] = self.chem_to_nut_result.get(1.0, tk.END)
        elif tab_index == 1:  # Nutrient to Chemical
            content['title'] = 'คำนวณสารเคมีจากธาตุอาหาร'
            content['data'] = self.nut_to_chem_result.get(1.0, tk.END)
        elif tab_index == 2:  # Formula Manager
            content['title'] = 'จัดการสูตร'
            content['data'] = self.formula_details.get(1.0, tk.END)
        elif tab_index == 3:  # Stock Solution
            content['title'] = 'สารละลายเข้มข้น'
            content['data'] = self.stock_result.get(1.0, tk.END)
        elif tab_index == 4:  # Database
            content['title'] = 'ฐานข้อมูลสารเคมี'
            content['data'] = self.db_text.get(1.0, tk.END)

        return content

    def export_to_pdf(self):
        """Export current tab content to PDF"""
        try:
            from tkinter import filedialog
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import inch

            content = self.get_current_tab_content()

            if not content['data'].strip():
                messagebox.showwarning("คำเตือน", "ไม่มีข้อมูลให้ส่งออก")
                return

            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )

            if filename:
                # Create PDF
                c = canvas.Canvas(filename, pagesize=A4)
                width, height = A4

                # Try to use Thai font if available
                try:
                    pdfmetrics.registerFont(TTFont('THSarabun', 'THSarabunNew.ttf'))
                    c.setFont('THSarabun', 16)
                except:
                    c.setFont('Helvetica', 12)

                # Title
                c.drawString(inch, height - inch, content['title'])

                # Content
                y_position = height - 1.5 * inch
                c.setFont('Courier', 9)

                for line in content['data'].split('\n'):
                    if y_position < inch:
                        c.showPage()
                        y_position = height - inch
                        c.setFont('Courier', 9)

                    # Handle Thai characters (basic approach)
                    try:
                        c.drawString(inch, y_position, line[:100])
                    except:
                        c.drawString(inch, y_position, line.encode('utf-8', errors='ignore').decode('utf-8')[:100])

                    y_position -= 12

                c.save()
                messagebox.showinfo("สำเร็จ", f"ส่งออกเป็น PDF เรียบร้อยแล้ว\n{filename}")

        except ImportError:
            messagebox.showerror("ข้อผิดพลาด",
                               "ไม่พบ reportlab library\nกรุณาติดตั้ง: pip install reportlab")
        except Exception as e:
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    def export_to_docx(self):
        """Export current tab content to DOCX"""
        try:
            from tkinter import filedialog
            from docx import Document
            from docx.shared import Pt

            content = self.get_current_tab_content()

            if not content['data'].strip():
                messagebox.showwarning("คำเตือน", "ไม่มีข้อมูลให้ส่งออก")
                return

            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
            )

            if filename:
                doc = Document()

                # Add title
                title = doc.add_heading(content['title'], 0)

                # Add content
                para = doc.add_paragraph()
                run = para.add_run(content['data'])
                run.font.size = Pt(10)
                run.font.name = 'Courier New'

                doc.save(filename)
                messagebox.showinfo("สำเร็จ", f"ส่งออกเป็น DOCX เรียบร้อยแล้ว\n{filename}")

        except ImportError:
            messagebox.showerror("ข้อผิดพลาด",
                               "ไม่พบ python-docx library\nกรุณาติดตั้ง: pip install python-docx")
        except Exception as e:
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    def export_to_csv(self):
        """Export current tab content to CSV"""
        try:
            from tkinter import filedialog
            import csv

            content = self.get_current_tab_content()

            if not content['data'].strip():
                messagebox.showwarning("คำเตือน", "ไม่มีข้อมูลให้ส่งออก")
                return

            filename = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
            )

            if filename:
                with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)

                    # Write title
                    writer.writerow([content['title']])
                    writer.writerow([])

                    # Write content line by line
                    for line in content['data'].split('\n'):
                        writer.writerow([line])

                messagebox.showinfo("สำเร็จ", f"ส่งออกเป็น CSV เรียบร้อยแล้ว\n{filename}")

        except Exception as e:
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")


def main():
    root = tk.Tk()
    app = NutrientCalculatorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
